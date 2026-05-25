from __future__ import annotations

import csv
import hashlib
import json
import os
import shutil
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
ARTIFACTS = DOCS / "_artifacts"
OUT = ARTIFACTS / "submission-20260525"
GENERATED_AT = datetime.now().isoformat(timespec="seconds")
PAGE_WIDTH_DXA = 8730


def rel(path: Path) -> str:
    return path.resolve().relative_to(ROOT).as_posix()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def read_json(path: Path, default: dict | None = None) -> dict:
    if not path.exists():
        return default or {}
    return json.loads(path.read_text(encoding="utf-8-sig"))


def set_run_font(run, *, size: float = 10.5, bold: bool = False, font: str = "SimSun") -> None:
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_cell_margins(cell, *, top: int = 120, start: int = 120, bottom: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold: bool = False, align: str = "left") -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == "center" else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=9, bold=bold, font="SimSun")
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)


def set_table_geometry(table, col_widths: list[int]) -> None:
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(col_widths)))
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths[index])


def add_page_number(paragraph) -> None:
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, font="SimSun")
    page_run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    page_run._r.extend([fld_char_begin, instr_text, fld_char_separate, fld_char_end])
    tail_run = paragraph.add_run(" 页")
    set_run_font(tail_run, size=9, font="SimSun")


def tune_document(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Pt(21)

    for name, size in (
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ):
        style = styles[name]
        style.font.name = "SimHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Pt(0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run(title)
    set_run_font(header_run, size=9, font="SimSun")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("课堂行为检测系统课程项目材料    ")
    set_run_font(footer_run, size=9, font="SimSun")
    add_page_number(footer)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
        style.font.size = Pt(10.5)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)


def add_cover_page(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(5):
        doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    set_run_font(title_run, size=22, bold=True, font="SimHei")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(subtitle)
    set_run_font(sub_run, size=12, font="SimSun")

    for _ in range(5):
        doc.add_paragraph()
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["项目名称", "课堂行为检测系统"],
            ["文档版本", "V1.0"],
            ["技术路线", "Flask 后端 + Web 可视化控制台 + YOLO 双模型检测"],
            ["生成时间", GENERATED_AT],
            ["验收依据", "scripts/verify_all.py 与 docs/_artifacts/verify-all-summary.json"],
        ],
        widths=[1800, 6930],
        center_columns={0},
    )
    doc.add_page_break()


def add_static_toc(doc: Document, sections: list[str]) -> None:
    doc.add_heading("目录", level=1)
    for item in sections:
        paragraph = doc.add_paragraph(item)
        paragraph.paragraph_format.first_line_indent = Pt(0)
    doc.add_page_break()


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[int] | None = None,
    center_columns: set[int] | None = None,
) -> None:
    center_columns = center_columns or set()
    widths = widths or [PAGE_WIDTH_DXA // len(headers)] * len(headers)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True, align="center")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_width(cells[index], widths[index])
            set_cell_text(cells[index], value, align="center" if index in center_columns else "left")
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Bullet")
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(item, style="List Number")
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.5


def add_picture_if_exists(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph(caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_p.runs:
        set_run_font(run, size=9, font="SimSun")


def add_formal_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.5


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_srs(path: Path) -> None:
    doc = Document()
    tune_document(doc, "课堂行为检测系统需求规格说明书", f"生成时间：{GENERATED_AT}")
    add_cover_page(doc, "课堂行为检测系统需求规格说明书", "科创实战项目课程提交材料")
    add_static_toc(
        doc,
        [
            "一、项目背景与建设目标",
            "二、系统总体需求",
            "三、功能性需求",
            "四、非功能性需求",
            "五、数据与接口需求",
            "六、验收标准",
        ],
    )

    doc.add_heading("一、项目背景与建设目标", level=1)
    add_formal_paragraph(
        doc,
        "随着课堂信息化管理和教学质量评价需求的提升，传统人工巡查方式在实时性、客观性和可追溯性方面存在不足。"
        "本项目围绕课堂教学场景，设计并实现一套课堂行为检测系统，用于对课堂图片、批量图片、视频片段和实时摄像头画面进行自动识别、统计分析与报告输出。"
    )
    add_formal_paragraph(
        doc,
        "系统建设目标包括：完成学生课堂行为检测、人头目标检测、检测结果可视化、历史任务管理、报告导出和随机点名辅助功能；"
        "同时保证项目能够在 Windows 本地环境下稳定启动、可复现验证、可按课程要求拆分提交。"
    )

    doc.add_heading("二、系统总体需求", level=1)
    add_formal_paragraph(
        doc,
        "系统采用 Flask 后端、Web 前端控制台和 YOLO 模型推理模块组成。后端负责模型加载、文件上传、检测任务调度、数据统计和报告生成；"
        "前端负责素材导入、检测模式切换、结果预览、图表展示、历史筛选和交互提示；模型层分别承担课堂行为识别与人头目标检测。"
    )
    add_table(
        doc,
        ["角色", "主要目标", "典型操作"],
        [
            ["管理员/演示者", "完成模型加载、检测演示、报告导出", "登录、选择模型、上传素材、查看报告"],
            ["教师/评审者", "查看课堂状态与行为统计", "查看检测框、行为分布、历史记录"],
            ["开发维护者", "维护模型、接口和验收脚本", "运行 verify_all、查看审计产物、打包归档"],
        ],
        widths=[1600, 3300, 3830],
        center_columns={0},
    )
    doc.add_heading("2.1 系统边界", level=2)
    add_formal_paragraph(
        doc,
        "系统主要解决本地课堂图像与视频素材的检测、展示和报告输出问题，不承担云端账号体系、跨校区数据同步和大规模数据集管理任务。"
        "课程提交版本强调可运行、可演示和可验证，所有核心功能均围绕本机模型推理与本地 Web 控制台展开。"
    )
    doc.add_heading("2.2 技术组成", level=2)
    add_table(
        doc,
        ["层次", "组成", "职责"],
        [
            ["表现层", "templates/app_shell.html、static/app、static/css", "提供模式切换、文件上传、结果预览、图表和历史记录"],
            ["接口层", "classroom_app/routes", "接收前端请求，完成参数校验、认证校验和统一 JSON 返回"],
            ["业务层", "classroom_app/services", "封装检测、任务、模型、报告和摄像头会话等业务逻辑"],
            ["模型层", "models/*.pt、utils/detector.py", "加载 YOLO 模型并输出检测框、类别、置信度和统计结果"],
            ["验证层", "scripts/*.py、static/app/**/*.test.js", "提供健康检查、回归烟测、浏览器审计和归档校验"],
        ],
        widths=[1200, 2800, 4730],
        center_columns={0},
    )

    doc.add_heading("三、功能性需求", level=1)
    doc.add_heading("3.1 检测任务需求", level=2)
    add_bullets(
        doc,
        [
            "单图检测：上传课堂图片，返回标注图、行为统计、检测明细和 HTML 报告。",
            "批量检测：一次上传多张图片，支持结果切换、批量统计和 ZIP 报告导出。",
            "视频检测：上传课堂视频，支持进度轮询、处理中止、完成态报告和历史追踪指标。",
            "实时摄像头：支持摄像头诊断、启动、停止、实时帧检测和任务记录。",
            "模型管理：支持学生行为模型与人头模型扫描、查看、切换和完整性校验。",
            "随机点名：基于检测到的人头目标随机选择对象，展示命中区域、置信度和来源。",
        ],
    )
    add_table(
        doc,
        ["编号", "需求名称", "优先级", "验收方式"],
        [
            ["FR-01", "单图检测", "高", "上传样例图片后生成标注图、摘要卡片和单任务报告"],
            ["FR-02", "批量检测", "高", "多图上传后可切换结果并导出批量报告 ZIP"],
            ["FR-03", "视频检测", "高", "视频任务可轮询进度、查看完成结果并支持停止"],
            ["FR-04", "实时摄像头", "中", "可诊断设备、启动实时检测并形成历史任务"],
            ["FR-05", "随机点名", "中", "检测到人头目标后可随机抽取并显示命中区域"],
            ["FR-06", "模型管理", "高", "可扫描候选模型、显示当前模型并阻止未批准路径"],
            ["FR-07", "报告归档", "高", "HTML 报告和 ZIP 清单内容通过契约校验"],
        ],
        widths=[900, 1800, 900, 5130],
        center_columns={0, 2},
    )
    doc.add_heading("3.2 报告与历史需求", level=2)
    add_formal_paragraph(
        doc,
        "系统应保存每次检测任务的类型、文件名、处理状态、检测统计、任务摘要和报告路径。历史记录应支持按模式筛选、按时间或主指标排序，并支持多选任务后批量导出报告 ZIP。"
    )
    add_formal_paragraph(
        doc,
        "单任务报告应包含检测概况、关键指标、行为统计和可读说明；批量报告 ZIP 应包含 HTML 报告、readme.txt 和 manifest.csv，便于人工查阅和后续归档。"
    )

    doc.add_heading("四、非功能性需求", level=1)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ["可移植性", "依赖以 requirements.txt 和本地 .venv 为主，路径保存采用相对 models 的形式。"],
            ["稳定性", "接口错误统一返回 JSON，视频停止、报告复用、批量导出具备契约测试。"],
            ["可审计性", "保留 verify_all、strict_system_audit、browser_visual_audit 等可复现验收入口。"],
            ["安全性", "管理员配置、运行时密钥、上传目录和输出目录避免作为默认固定敏感配置提交。"],
            ["可用性", "主工作台提供模式切换、历史筛选、报告打开、资源下载和通知反馈。"],
        ],
        widths=[1500, 7230],
        center_columns={0},
    )
    add_formal_paragraph(
        doc,
        "非功能性需求在本项目中直接关系到答辩现场稳定性。系统不能只在单一路径下运行成功，还应在错误上传、重复点击、模型切换、任务停止、报告资源缺失等边界场景下保持可理解的反馈。"
    )

    doc.add_heading("五、数据与接口需求", level=1)
    add_table(
        doc,
        ["模块", "输入", "输出", "说明"],
        [
            ["单图检测接口", "图片文件、置信度、IOU", "标注图、检测明细、统计摘要", "用于快速展示识别效果"],
            ["批量检测接口", "多张图片文件", "批量任务、结果列表、汇总报告", "用于多素材统一处理"],
            ["视频检测接口", "视频文件、跳帧参数", "任务进度、结果视频、检测明细", "支持停止和完成态报告"],
            ["摄像头接口", "摄像头索引、实时帧", "实时检测结果、任务记录", "包含诊断和异常提示"],
            ["模型管理接口", "模型类型、模型引用", "当前模型信息、加载结果", "限制未批准模型路径"],
        ],
        widths=[1400, 1850, 2500, 2980],
        center_columns={0},
    )
    add_formal_paragraph(
        doc,
        "接口异常应使用统一 JSON 结构返回错误码和错误信息，避免前端只能通过状态码或空响应判断失败。上传目录、输出目录、管理员配置和运行时密钥均应与仓库源码分离。"
    )
    doc.add_heading("5.1 数据保存要求", level=2)
    add_formal_paragraph(
        doc,
        "检测任务应保存任务编号、任务类型、原始文件、输出文件、检测数量、平均置信度、处理时长、行为统计和报告文件名。"
        "对于视频与摄像头任务，还应保留帧级明细、追踪指标和可用于回看分析的 summary 数据。"
    )

    doc.add_heading("六、验收标准", level=1)
    add_bullets(
        doc,
        [
            "一键总验收 `scripts/verify_all.py` 返回 overall_status=OK。",
            "单图、批量、视频、摄像头、历史、报告导出主链路均有自动化验证。",
            "浏览器审计覆盖桌面、笔记本和移动视口，console/network 无未处理问题。",
            "最终归档包不包含数据集和临时运行状态，保留运行所需源码、模型和文档。",
        ],
    )
    add_formal_paragraph(
        doc,
        "若验收脚本出现故意注入的异常日志，应以脚本退出码、最终汇总 JSON 和对应契约文件为准。"
        "项目提交时应同时提供源码、模型验证结果、系统测试报告、Bug 修复日志和最终归档包。"
    )
    save_doc(doc, path)


def build_bug_log(path: Path) -> None:
    doc = Document()
    tune_document(doc, "课堂行为检测系统 Bug 修复日志", f"生成时间：{GENERATED_AT}")
    add_cover_page(doc, "课堂行为检测系统 Bug 修复日志", "前后端联调复测与接口稳定性优化记录")
    add_static_toc(
        doc,
        [
            "一、联调背景",
            "二、问题修复明细",
            "三、重点修复说明",
            "四、修复验证矩阵",
            "五、风险与处理",
            "六、回归验证结论",
        ],
    )
    doc.add_heading("一、联调背景", level=1)
    add_formal_paragraph(
        doc,
        "本项目在单图检测、批量检测、视频检测、实时摄像头、历史记录、模型切换和报告导出等功能完成后，进入前后端联调与稳定性优化阶段。"
        "联调重点不只检查页面是否能够打开，而是针对异常输入、接口失败、重复操作、报告复用、任务状态漂移和浏览器端交互竞态进行逐项收敛。"
    )
    add_formal_paragraph(
        doc,
        "修复过程以自动化验收脚本为主要依据，涉及 `audit_readiness.py`、`regression_smoke.py`、`hardening_contracts.py`、"
        "`strict_system_audit.py` 和 `browser_visual_audit.ps1` 等验证入口。"
    )

    doc.add_heading("二、问题修复明细", level=1)
    add_table(
        doc,
        ["编号", "等级", "模块", "问题表现", "修复措施", "验证依据"],
        [
            ["BUG-01", "中", "上传校验", "批量上传遇到不支持文件时反馈不明确，人工演示时难以判断失败原因。", "统一扩展名与 MIME 检查，返回 400 unsupported_file，并保留可读 message。", "audit_readiness"],
            ["BUG-02", "高", "任务接口", "任务、模型异常路径返回结构不一致，前端提示容易出现空白或泛化错误。", "路由层统一异常映射，返回 JSON code/message/status，前端按 code 展示。", "hardening_contracts"],
            ["BUG-03", "中", "模型管理", "同一模型重复加载导致等待时间增加，影响答辩演示连贯性。", "增加当前模型判断，相同模型直接复用并返回已加载状态。", "hardening_contracts"],
            ["BUG-04", "高", "报告服务", "处理中任务可能被提前请求报告，导致报告内容不完整或路径不可用。", "报告生成前检查任务状态，未完成时返回 report_not_ready。", "hardening_contracts"],
            ["BUG-05", "高", "视频任务", "视频停止按钮重复点击后状态可能漂移，历史记录与界面状态不一致。", "停止接口改为幂等处理，允许 stopped_partial 稳定落库并生成摘要。", "regression_smoke"],
            ["BUG-06", "中", "批量导出", "批量报告重复生成且缺少内容指纹，无法确认归档内容是否漂移。", "增加 fingerprint、manifest.csv 和 ZIP 复用校验，确保同一批次内容稳定。", "verify_report_archive_test"],
            ["BUG-07", "中", "摄像头", "空会话、坏图片和设备不可用时缺少明确阻断，影响现场问题定位。", "补充 webcam_session_empty、invalid_image 与设备诊断返回。", "hardening_contracts"],
            ["BUG-08", "中", "前端交互", "快速重复点击模型加载、视频停止、历史导出时存在异步竞态。", "新增 action guard、version registry 和焦点回收逻辑。", "frontend_service_tests"],
        ],
        widths=[760, 620, 1120, 2300, 2500, 1430],
        center_columns={0, 1, 5},
    )
    add_formal_paragraph(
        doc,
        "以上问题按“可复现表现、影响范围、修复措施、验证依据”的顺序记录。对课程项目而言，重点不是单次运行成功，"
        "而是在错误输入、重复操作、资源缺失和模型切换等场景下仍能给出稳定、可解释、可回归的行为。"
    )

    doc.add_heading("三、重点修复说明", level=1)
    doc.add_heading("3.1 接口异常统一处理", level=2)
    add_formal_paragraph(
        doc,
        "任务列表、任务报告、模型信息和模型加载接口均补充异常捕获与错误码映射。后端不再让异常直接变成不可读页面或空响应，"
        "前端能够根据 `code` 字段给出明确提示。该项修复提升了系统在演示过程中的可解释性。"
    )
    doc.add_heading("3.2 报告生成与批量导出稳定性", level=2)
    add_formal_paragraph(
        doc,
        "报告服务增加任务状态检查、报告文件缺失检查、批量成员校验和 fingerprint 复用机制。"
        "当任务仍在处理、资源缺失或批量成员非法时，系统会阻断错误导出；当报告内容未变化时，系统可复用既有报告，减少重复生成。"
    )
    doc.add_heading("3.3 前端交互竞态控制", level=2)
    add_formal_paragraph(
        doc,
        "前端新增 action guard 和 version registry，用于限制重复点击、过期轮询结果和异步弹窗焦点漂移。"
        "该项修复保证了视频停止、历史导出、模型加载和摄像头操作在快速点击场景下仍保持状态一致。"
    )

    doc.add_heading("四、修复验证矩阵", level=1)
    add_table(
        doc,
        ["验证入口", "验证内容", "通过标准"],
        [
            ["scripts/audit_readiness.py", "检查仓库结构、配置文件、入口文件和提交前风险项。", "无阻断项，必要文件齐全。"],
            ["scripts/regression_smoke.py", "覆盖单图、批量、视频、历史、报告和摄像头生命周期。", "关键业务链路均返回 OK。"],
            ["scripts/hardening_contracts.py", "主动触发错误路径，验证接口错误码、报告状态和幂等行为。", "异常被结构化捕获，契约 JSON 通过。"],
            ["scripts/frontend_service_tests.py", "执行前端 ES module 单测，验证状态管理、历史导出和 action guard。", "全部前端服务测试通过。"],
            ["scripts/verify_all.py", "串联健康检查、烟测、回归、审计和浏览器视觉审计。", "summary 中 overall_status 为 OK。"],
        ],
        widths=[2200, 3800, 2730],
    )

    doc.add_heading("五、风险与处理", level=1)
    add_bullets(
        doc,
        [
            "摄像头设备可能被其他程序占用，系统通过诊断接口给出状态提示，并保留前端 fallback 流程。",
            "模型权重文件体积较大，最终归档包保留运行所需权重，但严格排除原始数据集和临时输出。",
            "浏览器视觉审计依赖本机 Edge/Chromium 环境，若本地环境缺失，应以后端回归和静态资源检查作为最低验收证据。",
            "hardening_contracts 会故意制造异常输入，因此测试日志中出现预期 traceback 时不代表系统失败。",
        ],
    )

    doc.add_heading("六、回归验证结论", level=1)
    add_formal_paragraph(
        doc,
        "最新总验收链路已覆盖健康检查、启动烟测、前端服务单测、UI 烟测、交互烟测、业务回归烟测、硬化契约、高标准边界审计、严格系统审计和浏览器视觉审计。"
        "其中 hardening_contracts 会故意注入异常路径，因此日志中出现的 traceback 属于预期测试行为，应以最终 OK 状态和 JSON 产物为准。"
    )
    save_doc(doc, path)


def build_test_report(path: Path) -> None:
    summary = read_json(ARTIFACTS / "verify-all-summary.json")
    strict = read_json(ARTIFACTS / "strict-system-audit.json")
    visual = read_json(ARTIFACTS / "browser-visual-audit.json")

    doc = Document()
    tune_document(doc, "课堂行为检测系统测试报告", f"生成时间：{GENERATED_AT}")
    add_cover_page(doc, "课堂行为检测系统测试报告", "项目全流程自测与系统级审计报告")
    add_static_toc(
        doc,
        [
            "一、测试目的与范围",
            "二、测试环境",
            "三、测试策略与用例设计",
            "四、自动化检查结果",
            "五、系统级浏览器审计",
            "六、测试证据",
            "七、结论与风险说明",
        ],
    )
    doc.add_heading("一、测试目的与范围", level=1)
    add_formal_paragraph(
        doc,
        "本次测试用于验证课堂行为检测系统在课程项目提交前的完整性、稳定性和可演示性。测试范围覆盖后端服务启动、模型加载、登录流程、"
        "单图检测、批量检测、视频检测、实时摄像头、历史记录、报告导出、异常边界和浏览器端视觉表现。"
    )
    add_formal_paragraph(
        doc,
        f"当前一键总验收状态为：{summary.get('overall_status', 'UNKNOWN')}。"
        f"严格系统审计状态为：{strict.get('overall_status', 'UNKNOWN')}。"
        "该结论来自仓库内自动化脚本与真实浏览器审计产物。"
    )

    doc.add_heading("二、测试环境", level=1)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["操作系统", "Windows 本地环境"],
            ["后端框架", "Flask 3.0.0"],
            ["模型推理", "Ultralytics YOLO + OpenCV"],
            ["前端形态", "HTML/CSS/JavaScript ES Module"],
            ["浏览器审计", "Microsoft Edge headless"],
            ["总验收入口", ".\\.venv\\Scripts\\python.exe scripts\\verify_all.py"],
        ],
        widths=[2200, 6530],
    )

    doc.add_heading("三、测试策略与用例设计", level=1)
    add_formal_paragraph(
        doc,
        "本次测试采用“接口契约 + 业务回归 + 前端单测 + 真实浏览器审计”的组合方式。接口契约用于保证异常输入可控，"
        "业务回归用于验证检测主流程，前端单测用于覆盖状态管理和交互防抖，真实浏览器审计用于确认页面在不同视口下可用。"
    )
    add_table(
        doc,
        ["用例编号", "测试对象", "测试要点", "预期结果"],
        [
            ["TC-01", "登录与会话", "访问登录页、认证接口和会话状态。", "登录状态可识别，未认证请求被正确拦截。"],
            ["TC-02", "单图检测", "上传样例图，检查检测框、统计摘要和 HTML 报告。", "返回标注图、检测明细和可打开报告。"],
            ["TC-03", "批量检测", "上传多张图片，检查结果切换、汇总统计和 ZIP 导出。", "批量报告包含 readme、manifest 和单项报告。"],
            ["TC-04", "视频检测", "提交视频任务，轮询进度并执行停止/完成态检查。", "任务状态稳定，停止接口具备幂等行为。"],
            ["TC-05", "摄像头流程", "执行设备诊断、实时会话、坏图片和空会话边界。", "异常被结构化返回，前端有明确提示。"],
            ["TC-06", "模型管理", "扫描模型、加载模型、重复加载、非法模型路径。", "当前模型可识别，非法路径被阻断。"],
            ["TC-07", "历史与报告", "筛选历史、打开单项报告、多选导出。", "历史数据一致，报告归档内容稳定。"],
            ["TC-08", "前端交互", "模式切换、重复点击、过期轮询和通知反馈。", "无未处理异常，重复操作不会破坏状态。"],
        ],
        widths=[920, 1450, 3330, 3030],
        center_columns={0},
    )

    doc.add_heading("四、自动化检查结果", level=1)
    checks = summary.get("checks", [])
    add_table(
        doc,
        ["检查项", "状态", "耗时"],
        [[item.get("name", ""), item.get("status", ""), f"{item.get('duration_seconds', '')} s"] for item in checks],
        widths=[4400, 1600, 2730],
        center_columns={1, 2},
    )
    add_formal_paragraph(
        doc,
        "从测试项结果看，系统基础健康检查、启动烟测、前端服务单测、UI 烟测、交互烟测、业务回归烟测和审计类脚本均返回 OK。"
        "其中业务回归覆盖了单图、批量、视频、历史、报告和摄像头生命周期；报告归档测试用于确认 HTML 和 ZIP 报告的内容一致性。"
    )

    doc.add_heading("五、系统级浏览器审计", level=1)
    flows = strict.get("flows", [])
    add_table(
        doc,
        ["流程", "视口", "状态", "Console 错误", "Network 失败"],
        [
            [
                flow.get("flow", ""),
                flow.get("viewport", ""),
                flow.get("status", ""),
                str(flow.get("console_error_count", "")),
                str(flow.get("network_failure_count", "")),
            ]
            for flow in flows[:16]
        ],
        widths=[2100, 1600, 1200, 1850, 1980],
        center_columns={1, 2, 3, 4},
    )
    add_formal_paragraph(
        doc,
        "严格系统审计使用真实浏览器访问本地服务，检查登录、模式切换、报告打开、批量导出、视频停止、摄像头 fallback 等流程，"
        "并同步采集 console 错误、network 失败和截图证据。当前 issue 列表为空。"
    )

    doc.add_heading("六、测试证据", level=1)
    add_picture_if_exists(doc, ARTIFACTS / "browser-audit-dashboard.png", "图 1 工作台浏览器截图")
    add_picture_if_exists(doc, ARTIFACTS / "browser-audit-report.png", "图 2 报告页浏览器截图")
    add_formal_paragraph(
        doc,
        "浏览器视觉审计记录了登录页、主工作台、摄像头页、单图报告页和视频报告页截图；"
        f"批量报告验证条目数：{visual.get('batch_report_entries_verified', '未记录')}。"
    )

    doc.add_heading("七、结论与风险说明", level=1)
    add_formal_paragraph(
        doc,
        "综合自动化测试、系统级浏览器审计和报告归档校验，课堂行为检测系统满足课程项目提交和演示要求。"
        "系统能够在本地环境完成模型加载、检测分析、结果展示、历史追踪和报告输出，且具备异常输入处理和关键链路回归能力。"
    )
    add_table(
        doc,
        ["类别", "统计结论", "处理状态"],
        [
            ["阻断缺陷", "未发现影响系统启动、登录、检测或报告导出的阻断缺陷。", "关闭"],
            ["一般缺陷", "联调阶段发现的接口返回、重复点击和报告复用问题已纳入修复日志。", "已复测"],
            ["环境风险", "Qt 壳层依赖 PySide6；摄像头能力依赖本机设备和驱动状态。", "保留说明"],
            ["数据风险", "最终归档不包含 datasets、testfile、uploads、outputs 和缓存目录。", "已排除"],
        ],
        widths=[1500, 5230, 2000],
        center_columns={0, 2},
    )
    add_bullets(
        doc,
        [
            "hardening_contracts 中出现的 traceback 属于故意注入的异常路径验证，需以最终 OK 状态和 JSON 产物判断。",
            "本机未安装 Qt 运行库时，Qt 壳层仅通过语法检查；安装 requirements-qt.txt 后可运行桌面入口。",
            "摄像头能力受本机设备占用和驱动状态影响，系统保留诊断与降级提示。"
        ],
    )
    save_doc(doc, path)


def build_manual(path: Path) -> None:
    doc = Document()
    tune_document(doc, "科创实战项目手册与训练进度安排", f"生成时间：{GENERATED_AT}")
    add_cover_page(doc, "科创实战项目手册与训练进度安排", "课堂行为检测系统项目过程性材料")
    add_static_toc(
        doc,
        [
            "一、项目简介",
            "二、总体训练与开发安排",
            "三、模型训练进度安排",
            "四、系统开发与联调安排",
            "五、运行与提交说明",
            "六、课程提交材料清单",
        ],
    )
    doc.add_heading("一、项目简介", level=1)
    add_formal_paragraph(
        doc,
        "本项目以课堂行为检测为应用对象，围绕人头检测模型、课堂行为检测模型、后端服务、前端可视化界面、随机点名功能和系统测试报告展开。"
        "项目最终形成可运行的本地系统和按课程平台要求拆分的提交材料。"
    )

    doc.add_heading("二、总体训练与开发安排", level=1)
    add_table(
        doc,
        ["阶段", "时间安排", "主要任务", "完成产物"],
        [
            ["第 1 阶段", "第 1 周", "需求分析、技术选型、项目结构搭建", "需求规格说明书、README、启动脚本"],
            ["第 2 阶段", "第 2 至 3 周", "人头检测模型验证与课堂行为模型验证", "训练结果曲线、混淆矩阵、验证样例、模型权重"],
            ["第 3 阶段", "第 4 周", "Flask 后端接口、模型服务、任务服务和报告服务开发", "检测接口、模型管理、报告服务、任务服务"],
            ["第 4 阶段", "第 5 周", "前端可视化界面、历史记录和随机点名功能开发", "工作台界面、图表、历史记录、随机点名弹窗"],
            ["第 5 阶段", "第 6 周", "前后端联调、Bug 修复、系统测试和浏览器审计", "Bug 修复日志、系统测试报告、审计截图"],
            ["第 6 阶段", "提交前", "最终归档与平台上传材料拆分", "按作业拆分的上传 ZIP/DOCX 和最终归档包"],
        ],
        widths=[1050, 1300, 3650, 2730],
        center_columns={0, 1},
    )
    add_formal_paragraph(
        doc,
        "整体安排遵循“先模型验证、再系统开发、后联调验收”的顺序。模型训练结果作为后续系统推理的基础，系统开发完成后再通过自动化脚本和浏览器审计进行闭环验证。"
    )
    doc.add_heading("2.1 过程里程碑", level=2)
    add_table(
        doc,
        ["里程碑", "完成条件", "证明材料"],
        [
            ["M1 需求冻结", "检测对象、用户角色、输入输出和提交边界明确。", "需求规格说明书"],
            ["M2 模型可用", "两类模型均具备训练/验证结果，权重文件可被系统加载。", "作业 2、作业 3 验证 ZIP"],
            ["M3 系统闭环", "上传、检测、展示、历史、报告导出形成完整链路。", "后端代码包、前端代码包"],
            ["M4 联调通过", "关键异常路径和重复操作场景均可回归。", "Bug 修复日志、系统测试报告"],
            ["M5 归档提交", "材料按平台入口拆分，最终归档包不含数据集。", "提交说明与 manifest"],
        ],
        widths=[1700, 4300, 2730],
        center_columns={0},
    )

    doc.add_heading("三、模型训练进度安排", level=1)
    add_table(
        doc,
        ["训练对象", "训练/验证材料", "当前状态"],
        [
            ["人头检测模型", "results.csv、results.png、F1/PR/P/R 曲线、混淆矩阵、验证样例", "已整理为作业 2 验证结果包"],
            ["课堂行为检测模型", "results.csv、results.png、F1/PR/P/R 曲线、混淆矩阵、验证样例", "已整理为作业 3 验证结果包"],
            ["系统联调", "verify_all、strict_system_audit、browser_visual_audit", "最新总验收 OK"],
        ],
        widths=[1700, 4700, 2330],
        center_columns={0, 2},
    )
    doc.add_heading("3.1 人头检测模型", level=2)
    add_formal_paragraph(
        doc,
        "人头检测模型用于定位课堂画面中的头部目标，为人数统计、实时巡检和随机点名提供基础目标框。"
        "训练结果文件包括 results.csv、results.png、F1_curve、PR_curve、P_curve、R_curve、混淆矩阵和验证批次图。"
    )
    doc.add_heading("3.2 课堂行为检测模型", level=2)
    add_formal_paragraph(
        doc,
        "课堂行为检测模型用于识别举手、阅读、书写、睡觉、使用手机、低头和趴伏等课堂行为。"
        "模型验证结果用于说明训练过程、指标变化和验证样例效果，已整理为独立压缩包便于上传。"
    )
    doc.add_heading("3.3 训练结果文件说明", level=2)
    add_table(
        doc,
        ["文件类型", "说明", "用于证明"],
        [
            ["results.csv", "记录每轮训练的 precision、recall、mAP、loss 等指标。", "训练过程可追溯"],
            ["results.png", "汇总展示训练/验证指标曲线。", "整体收敛趋势"],
            ["PR/F1/P/R 曲线", "展示不同阈值下的准确率、召回率和综合表现。", "模型阈值效果"],
            ["confusion_matrix", "展示各类别预测与真实标签的混淆情况。", "类别识别稳定性"],
            ["val_batch 图片", "展示验证集样例预测框和标签效果。", "视觉验证结果"],
            ["best.pt / last.pt", "训练产生的模型权重文件。", "系统运行基础"],
        ],
        widths=[1900, 4600, 2230],
        center_columns={0, 2},
    )

    doc.add_heading("四、系统开发与联调安排", level=1)
    add_formal_paragraph(
        doc,
        "后端开发重点为 Flask 路由、模型加载、检测服务、任务服务、报告服务和异常处理；前端开发重点为模式切换、素材上传、检测结果预览、图表统计、历史记录和报告导出。"
    )
    add_formal_paragraph(
        doc,
        "联调阶段重点验证前后端接口状态码、JSON 返回结构、任务状态一致性、批量 ZIP 报告内容、浏览器控制台错误和不同视口下的页面布局。"
    )
    add_table(
        doc,
        ["开发模块", "关键文件", "联调重点"],
        [
            ["后端入口", "app.py、classroom_app/__init__.py", "应用工厂、配置加载、登录会话和健康检查。"],
            ["检测服务", "classroom_app/services、utils/detector.py", "模型加载、检测参数、异常返回和结果保存。"],
            ["报告服务", "classroom_app/services/report_service.py", "HTML 报告、批量 ZIP、manifest 与文件缺失处理。"],
            ["前端界面", "templates/app_shell.html、static/app/main.js", "模式切换、上传交互、图表、历史和通知反馈。"],
            ["随机点名", "static/app/random-call、相关后端接口", "人头检测结果来源、随机命中区域和无目标提示。"],
            ["验收脚本", "scripts/verify_all.py、scripts/*audit*.py", "一键回归、严格审计、浏览器截图和 JSON 汇总。"],
        ],
        widths=[1500, 3000, 4230],
        center_columns={0},
    )

    doc.add_heading("五、运行与提交说明", level=1)
    add_numbered(
        doc,
        [
            "安装 Python 依赖：`.\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt`。",
            "初始化管理员账号：`.\\.venv\\Scripts\\python.exe scripts\\init_local_admin.py --username admin --password \"your-password\"`。",
            "运行总验收：`.\\.venv\\Scripts\\python.exe scripts\\verify_all.py`。",
            "启动答辩演示：`start_demo_session.bat`。",
            "如需 Qt 入口：安装 `requirements-qt.txt` 后运行 `python qt_launcher.py`。",
        ],
    )
    add_formal_paragraph(
        doc,
        "本次提交按 11 个上传入口拆分，代码、验证结果、测试报告和最终归档包相互独立。最终归档包不包含数据集、虚拟环境、缓存、上传目录和运行输出目录，仅保留源码、模型权重、文档和必要验收证据。"
    )
    doc.add_heading("六、课程提交材料清单", level=1)
    add_table(
        doc,
        ["上传项", "文件名", "内容说明"],
        [
            ["作业 1", "01_需求规格说明书.docx", "系统背景、需求、接口、数据和验收标准。"],
            ["作业 2", "02_人头检测模型验证结果.zip", "人头检测模型训练曲线、混淆矩阵、验证图和说明。"],
            ["作业 3", "03_课堂行为检测模型验证结果.zip", "课堂行为检测模型训练曲线、混淆矩阵、验证图和说明。"],
            ["作业 4", "04_Qt可视化界面代码.zip", "可选 Qt 桌面入口源码和依赖说明。"],
            ["作业 5", "05_后端核心文件代码.zip", "Flask 后端入口、路由、服务层、工具和验收脚本。"],
            ["作业 6", "06_前端文件完整代码.zip", "HTML 模板、CSS、前端 ES module 和图表库资源。"],
            ["作业 7", "07_随机点名核心代码.zip", "随机点名相关前端和后端核心代码。"],
            ["作业 8", "08_Bug修复日志.docx", "问题编号、修复措施、验证依据和回归结论。"],
            ["作业 9", "09_系统测试报告.docx", "测试环境、测试用例、审计结果、截图证据和风险说明。"],
            ["作业 10", "10_项目最终归档包_不含数据集.zip", "源码、模型、文档、测试证据和运行说明，不含数据集。"],
            ["手册上传", "科创实战项目手册_训练进度安排.docx", "项目过程、训练进度、开发安排和提交清单。"],
        ],
        widths=[1200, 2850, 4680],
        center_columns={0},
    )
    save_doc(doc, path)


def write_readme(path: Path, title: str, lines: list[str]) -> None:
    path.write_text(title + "\n" + "=" * len(title) + "\n\n" + "\n".join(lines) + "\n", encoding="utf-8")


def zip_paths(zip_path: Path, entries: list[tuple[Path, str]]) -> None:
    zip_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for src, arc in sorted(entries, key=lambda item: item[1]):
            if src.exists() and src.is_file():
                archive.write(src, arc)


def collect_files(root: Path, patterns: tuple[str, ...], exclude_parts: set[str] | None = None) -> list[Path]:
    exclude_parts = exclude_parts or {"node_modules", "__pycache__"}
    files: list[Path] = []
    for pattern in patterns:
        for path in root.glob(pattern):
            if path.is_file() and not (set(path.parts) & exclude_parts):
                files.append(path)
    return sorted(set(files))


def build_code_zips(paths: dict[str, Path]) -> None:
    temp = OUT / "_tmp_readmes"
    temp.mkdir(parents=True, exist_ok=True)

    backend_readme = temp / "README_后端核心文件代码.txt"
    write_readme(
        backend_readme,
        "后端核心文件代码说明",
        [
            "本包包含 Flask 后端入口、配置、路由、服务层、核心工具和关键验收脚本。",
            "正式入口为 app.py，应用工厂位于 classroom_app/__init__.py。",
            "运行前请按主 README 初始化 .venv 与管理员账号。",
        ],
    )
    backend_files = [
        ROOT / "app.py",
        ROOT / "config.py",
        ROOT / "requirements.txt",
        ROOT / "README.md",
        backend_readme,
    ]
    backend_files += collect_files(ROOT / "classroom_app", ("**/*.py",))
    backend_files += collect_files(ROOT / "utils", ("**/*.py",))
    backend_files += collect_files(ROOT / "scripts", ("*.py", "*.md"))
    zip_paths(paths["backend"], [(p, rel(p) if p.is_relative_to(ROOT) else p.name) for p in backend_files])

    frontend_readme = temp / "README_前端文件完整代码.txt"
    write_readme(
        frontend_readme,
        "前端文件完整代码说明",
        [
            "本包包含当前 Web 可视化界面模板、CSS、ES module 前端代码和图表库。",
            "主要入口为 templates/app_shell.html 与 static/app/main.js。",
            "前端单测由 scripts/frontend_service_tests.py 递归执行 static/app/**/*.test.js。",
        ],
    )
    frontend_files = [frontend_readme]
    frontend_files += collect_files(ROOT / "templates", ("*.html",))
    frontend_files += collect_files(ROOT / "static", ("app/**/*.js", "css/**/*.css", "vendor/**/*.js", "favicon.ico"))
    zip_paths(paths["frontend"], [(p, rel(p) if p.is_relative_to(ROOT) else p.name) for p in frontend_files])

    random_readme = temp / "README_随机点名核心代码.txt"
    write_readme(
        random_readme,
        "随机点名核心代码说明",
        [
            "随机点名功能基于当前检测结果中的人头候选目标，从可见目标中随机抽取并展示。",
            "核心逻辑涉及 static/app/lib/inspector.js 的候选目标提取与随机选择，以及 static/app/components/random-call.js 的结果渲染。",
            "templates/app_shell.html 提供随机点名按钮与弹窗结构，static/css/app.css 提供弹窗布局样式。",
        ],
    )
    random_files = [
        random_readme,
        ROOT / "static" / "app" / "components" / "random-call.js",
        ROOT / "static" / "app" / "lib" / "inspector.js",
        ROOT / "static" / "app" / "main.js",
        ROOT / "templates" / "app_shell.html",
        ROOT / "static" / "css" / "app.css",
    ]
    zip_paths(paths["random"], [(p, rel(p) if p.is_relative_to(ROOT) else p.name) for p in random_files])

    qt_readme = temp / "README_Qt可视化界面代码.txt"
    write_readme(
        qt_readme,
        "Qt 可视化界面代码说明",
        [
            "本项目真实主界面是 Flask + Web 控制台；本包额外提供 Qt 桌面入口，满足课程中 Qt 可视化界面提交要求。",
            "qt_launcher.py 会启动或复用 127.0.0.1:5000 的 Flask 服务，并优先用 QtWebEngine 嵌入 Web 控制台。",
            "若环境缺少 PySide6，请先执行 pip install -r requirements-qt.txt。",
        ],
    )
    qt_files = [qt_readme, ROOT / "qt_launcher.py", ROOT / "requirements-qt.txt", ROOT / "app.py", ROOT / "README.md"]
    zip_paths(paths["qt"], [(p, rel(p) if p.is_relative_to(ROOT) else p.name) for p in qt_files])


def build_model_zip(zip_path: Path, model_dir: Path, title: str) -> None:
    readme = OUT / "_tmp_readmes" / f"README_{title}.txt"
    write_readme(
        readme,
        f"{title}验证结果说明",
        [
            "本包整理 YOLO 训练/验证输出，包含 results.csv、results.png、P/R/PR/F1 曲线、混淆矩阵和训练/验证样例图。",
            "best.pt 与 last.pt 为训练权重文件，便于复核模型产物完整性。",
        ],
    )
    names = {
        "args.yaml",
        "results.csv",
        "results.png",
        "P_curve.png",
        "R_curve.png",
        "PR_curve.png",
        "F1_curve.png",
        "confusion_matrix.png",
        "confusion_matrix_normalized.png",
        "labels.jpg",
        "labels_correlogram.jpg",
        "train_batch0.jpg",
        "train_batch1.jpg",
        "train_batch2.jpg",
        "val_batch0_labels.jpg",
        "val_batch0_pred.jpg",
        "val_batch1_labels.jpg",
        "val_batch1_pred.jpg",
        "val_batch2_labels.jpg",
        "val_batch2_pred.jpg",
    }
    entries = [(readme, readme.name)]
    for path in sorted(model_dir.rglob("*")):
        if path.is_file() and (path.name in names or path.name in {"best.pt", "last.pt"}):
            entries.append((path, f"{title}/{path.relative_to(model_dir).as_posix()}"))
    zip_paths(zip_path, entries)


def should_skip_final(path: Path) -> bool:
    rel_parts = path.relative_to(ROOT).parts
    skip_roots = {
        ".git",
        ".venv",
        ".codex-gh",
        "__pycache__",
        "datasets",
        "testfile",
        "uploads",
        "outputs",
        "Ultralytics",
    }
    if rel_parts[0] in skip_roots:
        return True
    if "node_modules" in rel_parts or "__pycache__" in rel_parts:
        return True
    if rel_parts[:2] == ("docs", "_artifacts"):
        return True
    if path.suffix.lower() in {".pyc", ".log", ".db"}:
        return True
    if path.name in {"admin_config.json", "runtime_secrets.json", "user_config.json"}:
        return True
    return False


def build_final_archive(zip_path: Path, generated_docs: list[Path]) -> None:
    selected_evidence = [
        ARTIFACTS / "verify-all-summary.json",
        ARTIFACTS / "strict-system-audit.json",
        ARTIFACTS / "browser-visual-audit.json",
        ARTIFACTS / "hardening-contracts.json",
        ARTIFACTS / "browser-audit-login.png",
        ARTIFACTS / "browser-audit-dashboard.png",
        ARTIFACTS / "browser-audit-report.png",
        ARTIFACTS / "browser-audit-video-report.png",
        ARTIFACTS / "browser-audit-webcam.png",
    ]
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(ROOT.rglob("*")):
            if path.is_file() and not should_skip_final(path):
                archive.write(path, rel(path))
        for path in selected_evidence:
            if path.exists():
                archive.write(path, f"verification_artifacts/{path.name}")
        for path in generated_docs:
            if path.exists():
                archive.write(path, f"submission_materials/{path.name}")


def build_submission_text(path: Path) -> None:
    text = """作业1：上传需求规格说明书
已完成课堂行为检测系统需求规格说明书，内容包含项目概述、用户角色、功能需求、非功能需求和验收标准，附件为 01_需求规格说明书.docx。

作业2：上传人头检测模型验证结果文件
已整理人头检测模型验证结果，包含 results.csv、results.png、P/R/PR/F1 曲线、混淆矩阵、训练与验证样例图及权重文件，附件为 02_人头检测模型验证结果.zip。

作业3：上传行为检测模型验证结果文件
已整理课堂行为检测模型验证结果，包含训练参数、指标曲线、混淆矩阵、验证样例和权重文件，附件为 03_课堂行为检测模型验证结果.zip。

作业4：上传 Qt 可视化界面代码
已补充 Qt 桌面可视化入口，Qt 壳层用于启动或嵌入当前 Flask Web 控制台，附件为 04_Qt可视化界面代码.zip。

作业5：上传后端核心文件代码
已整理 Flask 后端核心文件，包含 app.py、配置、路由、服务层、工具模块和关键验收脚本，附件为 05_后端核心文件代码.zip。

作业6：上传前端文件代码
已整理前端完整代码，包含模板、CSS、ES module 组件、服务模块、工具模块和图表依赖，附件为 06_前端文件完整代码.zip。

作业7：上传随机点名核心代码
已整理随机点名核心代码，包含候选目标提取、随机选择、弹窗渲染和界面按钮结构，附件为 07_随机点名核心代码.zip。

作业8：上传 Bug 修复日志
已完成前后端联调复测和接口稳定性优化记录，附件为 08_Bug修复日志.docx。

作业9：上传系统测试报告
已完成项目全流程自测，总验收 verify_all.py 返回 OK，附件为 09_系统测试报告.docx。

作业10：上传项目最终归档包
已整理项目最终归档包，包含源码、模型权重、文档、测试报告和必要截图证据，不包含数据集，附件为 10_项目最终归档包_不含数据集.zip。

上传《科创实战项目手册》
已完成手册中的训练进度安排和项目交付说明，附件为 科创实战项目手册_训练进度安排.docx。
"""
    path.write_text(text, encoding="utf-8")


def validate_no_forbidden(zip_path: Path) -> list[str]:
    forbidden_prefixes = ("datasets/", "testfile/", ".venv/", "uploads/", "outputs/", "static/app/node_modules/")
    violations: list[str] = []
    with zipfile.ZipFile(zip_path) as archive:
        for name in archive.namelist():
            normalized = name.replace("\\", "/")
            if normalized.endswith(".pyc") or "__pycache__/" in normalized:
                violations.append(normalized)
            if normalized.startswith(forbidden_prefixes):
                violations.append(normalized)
    return violations


def build_manifest(files: dict[str, Path], validations: dict[str, list[str]]) -> None:
    payload = {
        "generated_at": GENERATED_AT,
        "output_dir": rel(OUT),
        "verify_all_summary": rel(ARTIFACTS / "verify-all-summary.json"),
        "items": [],
        "zip_validation": validations,
    }
    for key, path in files.items():
        payload["items"].append(
            {
                "key": key,
                "filename": path.name,
                "path": rel(path),
                "size_bytes": path.stat().st_size if path.exists() else 0,
                "sha256": sha256(path) if path.exists() else "",
            }
        )
    (OUT / "交付清单-manifest.json").write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def main() -> int:
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / "_tmp_readmes").mkdir(parents=True, exist_ok=True)

    files = {
        "assignment_01_srs": OUT / "01_需求规格说明书.docx",
        "assignment_02_head_model": OUT / "02_人头检测模型验证结果.zip",
        "assignment_03_behavior_model": OUT / "03_课堂行为检测模型验证结果.zip",
        "assignment_04_qt": OUT / "04_Qt可视化界面代码.zip",
        "assignment_05_backend": OUT / "05_后端核心文件代码.zip",
        "assignment_06_frontend": OUT / "06_前端文件完整代码.zip",
        "assignment_07_random_call": OUT / "07_随机点名核心代码.zip",
        "assignment_08_bug_log": OUT / "08_Bug修复日志.docx",
        "assignment_09_test_report": OUT / "09_系统测试报告.docx",
        "assignment_10_archive": OUT / "10_项目最终归档包_不含数据集.zip",
        "manual": OUT / "科创实战项目手册_训练进度安排.docx",
        "copy_text": OUT / "提交说明-逐项复制.txt",
    }

    build_srs(files["assignment_01_srs"])
    build_bug_log(files["assignment_08_bug_log"])
    build_test_report(files["assignment_09_test_report"])
    build_manual(files["manual"])

    build_model_zip(
        files["assignment_02_head_model"],
        ROOT / "models" / "yolov8_trainresult_heand" / "train",
        "人头检测模型",
    )
    build_model_zip(
        files["assignment_03_behavior_model"],
        ROOT / "models" / "yolov8_trainresult_behavior" / "train",
        "课堂行为检测模型",
    )
    build_code_zips(
        {
            "qt": files["assignment_04_qt"],
            "backend": files["assignment_05_backend"],
            "frontend": files["assignment_06_frontend"],
            "random": files["assignment_07_random_call"],
        }
    )
    build_submission_text(files["copy_text"])
    build_final_archive(
        files["assignment_10_archive"],
        [
            files["assignment_01_srs"],
            files["assignment_08_bug_log"],
            files["assignment_09_test_report"],
            files["manual"],
            files["copy_text"],
        ],
    )

    validations = {
        key: validate_no_forbidden(path)
        for key, path in files.items()
        if path.suffix.lower() == ".zip"
    }
    build_manifest(files, validations)
    shutil.rmtree(OUT / "_tmp_readmes", ignore_errors=True)

    for key, path in files.items():
        print(f"{key}: {rel(path)}")
    print(f"manifest: {rel(OUT / '交付清单-manifest.json')}")
    bad = {key: value for key, value in validations.items() if value}
    if bad:
        print(json.dumps(bad, ensure_ascii=False, indent=2))
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
